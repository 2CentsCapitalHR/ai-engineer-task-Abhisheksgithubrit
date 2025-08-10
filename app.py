import streamlit as st
import os
import json
import tempfile
import pandas as pd
from docx import Document
import re
import base64
from io import BytesIO

# Set page configuration
st.set_page_config(
    page_title="ADGM-Compliant Corporate Agent",
    page_icon="📝",
    layout="wide"
)

# Define document categories and their required documents
DOCUMENT_REQUIREMENTS = {
    "Company Formation": [
        "Articles of Association (AoA)",
        "Memorandum of Association (MoA/MoU)",
        "Board Resolution Templates",
        "Shareholder Resolution Templates",
        "Incorporation Application Form",
        "UBO Declaration Form",
        "Register of Members and Directors",
        "Change of Registered Address Notice"
    ],
    "Licensing Regulatory Filings": [
        "Application for Commercial License",
        "Regulatory Business Plan",
        "Business Activity Details",
        "Data Protection Notification Form",
        "Fitness and Propriety Form"
    ],
    "Employment HR Contracts": [
        "Employment Contract",
        "Employee Handbook",
        "Confidentiality Agreement",
        "Non-Disclosure Agreement",
        "Termination Letter Template"
    ],
    "Commercial Agreements": [
        "Service Level Agreement",
        "Non-Disclosure Agreement",
        "Joint Venture Agreement",
        "Supplier Agreement",
        "Client Contract Template"
    ],
    "Compliance Risk Policies": [
        "Anti-Money Laundering Policy",
        "Risk Management Framework",
        "Compliance Manual",
        "Data Protection Policy",
        "Business Continuity Plan"
    ]
}

# ADGM specific regulations and common issues
ADGM_REGULATIONS = {
    "Articles of Association": [
        {"pattern": r"UAE Federal (Court|law)", "issue": "Incorrect jurisdiction (should reference ADGM Courts)", "severity": "High", "regulation": "ADGM Companies Regulations 2020, Art. 4"},
        {"pattern": r"share capital(?!.*denominated.*in (USD|US dollars|United States dollars))", "issue": "Share capital must be denominated in USD", "severity": "High", "regulation": "ADGM Companies Regulations 2020, Art. 8(2)"},
        {"pattern": r"director.{1,20}(not less than|minimum|at least) (\d+)", "issue": "ADGM requires at least one director who is a natural person", "severity": "High", "regulation": "ADGM Companies Regulations 2020, Art. 143"}
    ],
    "Memorandum of Association": [
        {"pattern": r"Dubai|Sharjah|Ajman|Umm Al Quwain|Fujairah|Ras Al Khaimah", "issue": "Incorrect reference to other UAE Emirates", "severity": "Medium", "regulation": "ADGM Companies Regulations 2020"},
        {"pattern": r"UAE Ministry", "issue": "References to UAE Ministry may be incorrect as ADGM has its own regulatory framework", "severity": "Medium", "regulation": "ADGM Founding Law (Abu Dhabi Law No. 4 of 2013)"}
    ],
    "UBO Declaration Form": [
        {"pattern": r"beneficial owner.{1,50}25%", "issue": "ADGM threshold for UBO is 25% ownership or control", "severity": "High", "regulation": "ADGM UBO Regulations 2019"},
        {"pattern": r"shareholder.{1,30}legal.{1,30}beneficial", "issue": "Clear distinction between legal and beneficial ownership needed", "severity": "Medium", "regulation": "ADGM UBO Regulations 2019, Art. 6"}
    ],
    "Employment Contract": [
        {"pattern": r"UAE Labour Law", "issue": "ADGM has its own Employment Regulations", "severity": "High", "regulation": "ADGM Employment Regulations 2019"},
        {"pattern": r"probation.{1,20}(more than|exceeding|over|above).{1,5}6 month", "issue": "Probation period cannot exceed 6 months under ADGM Employment Regulations", "severity": "High", "regulation": "ADGM Employment Regulations 2019, Art. 8"}
    ],
    "Board Resolution": [
        {"pattern": r"approve.{1,50}dividend.{1,50}(?!solvency)", "issue": "Dividend declaration must include solvency statement", "severity": "High", "regulation": "ADGM Companies Regulations 2020, Art. 107"},
        {"pattern": r"electronic signature", "issue": "Validate electronic signature compliance with ADGM Electronic Transactions Regulations", "severity": "Medium", "regulation": "ADGM Electronic Transactions Regulations 2021"}
    ],
    "General": [
        {"pattern": r"(Dubai|Ajman|Sharjah|UAE).{1,20}court", "issue": "Incorrect jurisdiction (should be ADGM Courts)", "severity": "High", "regulation": "ADGM Courts Regulations 2015"},
        {"pattern": r"UAE dirham|AED|Dirham", "issue": "Currency should be USD for ADGM companies", "severity": "Medium", "regulation": "ADGM Commercial Licensing Regulations 2015, Art. 12"},
        {"pattern": r"(?<!Abu Dhabi )Global Market", "issue": "Incorrect reference to ADGM", "severity": "Low", "regulation": "ADGM Founding Law (Abu Dhabi Law No. 4 of 2013)"}
    ]
}

# Official ADGM Document Resources
ADGM_RESOURCES = [
    {
        "category": "Company Formation & Governance",
        "document_type": "General Incorporation, AoA, MoA, Registers, UBO, Board Resolutions",
        "link": "https://www.adgm.com/registration-authority/registration-and-incorporation"
    },
    {
        "category": "Company Formation",
        "document_type": "Resolution for Incorporation (LTD - Multiple Shareholders)",
        "link": "https://assets.adgm.com/download/assets/adgm-ra-resolution-multiple-incorporate-shareholders-LTD-incorporation-v2.docx/186a12846c3911efa4e6c6223862cd87"
    },
    {
        "category": "Company Formation & Compliance",
        "document_type": "Incorporation, SPV, LLC, Other Forms & Templates",
        "link": "https://www.adgm.com/setting-up"
    },
    {
        "category": "Policy & Guidance",
        "document_type": "Guidance, Templates, Policy Statements",
        "link": "https://www.adgm.com/legal-framework/guidance-and-policy-statements"
    },
    {
        "category": "ADGM Company Set-up",
        "document_type": "Checklist – Company Set-up (Various Entities)",
        "link": "https://www.adgm.com/documents/registration-authority/registration-and-incorporation/checklist/branch-non-financial-services-20231228.pdf"
    },
    {
        "category": "ADGM Company Set-up",
        "document_type": "Checklist – Private Company Limited",
        "link": "https://www.adgm.com/documents/registration-authority/registration-and-incorporation/checklist/private-company-limited-by-guarantee-non-financial-services-20231228.pdf"
    },
    {
        "category": "Employment & HR",
        "document_type": "Standard Employment Contract Template (2024 update)",
        "link": "https://assets.adgm.com/download/assets/ADGM+Standard+Employment+Contract+Template+-+ER+2024+(Feb+2025).docx/ee14b252edbe11efa63b12b3a30e5e3a"
    },
    {
        "category": "Employment & HR",
        "document_type": "Standard Employment Contract Template (2019 short version)",
        "link": "https://assets.adgm.com/download/assets/ADGM+Standard+Employment+Contract+-+ER+2019+-+Short+Version+(May+2024).docx/33b57a92ecfe11ef97a536cc36767ef8"
    },
    {
        "category": "Data Protection",
        "document_type": "Appropriate Policy Document Template",
        "link": "https://www.adgm.com/documents/office-of-data-protection/templates/adgm-dpr-2021-appropriate-policy-document.pdf"
    },
    {
        "category": "Compliance & Filings",
        "document_type": "Annual Accounts & Filings",
        "link": "https://www.adgm.com/operating-in-adgm/obligations-of-adgm-registered-entities/annual-filings/annual-accounts"
    },
    {
        "category": "Letters/Permits",
        "document_type": "Application for Official Letters & Permits",
        "link": "https://www.adgm.com/operating-in-adgm/post-registration-services/letters-and-permits"
    },
    {
        "category": "Regulatory Guidance",
        "document_type": "Incorporation Package, Filings, Templates",
        "link": "https://en.adgm.thomsonreuters.com/rulebook/7-company-incorporation-package"
    },
    {
        "category": "Regulatory Template",
        "document_type": "Shareholder Resolution – Amendment of Articles",
        "link": "https://assets.adgm.com/download/assets/Templates_SHReso_AmendmentArticles-v1-20220107.docx/97120d7c5af911efae4b1e183375c0b2?forcedownload=1"
    }
]

def create_document_type_classifier():
    """Create a simple keyword-based document classifier"""
    document_types = {}
    for category, docs in DOCUMENT_REQUIREMENTS.items():
        for doc in docs:
            # Extract main keywords from document names to use for classification
            keywords = re.findall(r'\b[A-Za-z]{3,}\b', doc)
            keywords = [k.lower() for k in keywords if k.lower() not in ['and', 'the', 'for', 'of']]
            document_types[doc] = {
                'keywords': keywords,
                'category': category
            }
    return document_types

def classify_document(text, classifier):
    """Classify document based on its content"""
    text_lower = text.lower()
    best_match = None
    highest_score = 0
    
    for doc_type, info in classifier.items():
        score = sum(keyword in text_lower for keyword in info['keywords'])
        if score > highest_score:
            highest_score = score
            best_match = doc_type
    
    return best_match if highest_score > 0 else None

def detect_red_flags(document_text, document_type):
    """Detect red flags in the document based on ADGM regulations"""
    issues = []
    
    # Check document-specific rules
    if document_type in ADGM_REGULATIONS:
        rules = ADGM_REGULATIONS[document_type]
    else:
        rules = ADGM_REGULATIONS.get("General", [])
    
    for rule in rules:
        matches = re.finditer(rule["pattern"], document_text, re.IGNORECASE)
        for match in matches:
            issues.append({
                "document": document_type,
                "section": f"Text near position {match.start()}",
                "issue": rule["issue"],
                "severity": rule["severity"],
                "suggestion": f"Align with {rule['regulation']}",
                "matched_text": match.group(0),
                "position": match.start()
            })
    
    # Also check general rules for all documents
    if document_type != "General":
        general_rules = ADGM_REGULATIONS.get("General", [])
        for rule in general_rules:
            matches = re.finditer(rule["pattern"], document_text, re.IGNORECASE)
            for match in matches:
                issues.append({
                    "document": document_type,
                    "section": f"Text near position {match.start()}",
                    "issue": rule["issue"],
                    "severity": rule["severity"],
                    "suggestion": f"Align with {rule['regulation']}",
                    "matched_text": match.group(0),
                    "position": match.start()
                })
    
    return issues

def insert_comments(doc_obj, issues):
    """Insert comments in the document at positions where issues were found"""
    # Sort issues by position in reverse order to avoid offset changes
    sorted_issues = sorted(issues, key=lambda x: x["position"], reverse=True)
    
    for issue in sorted_issues:
        position = issue["position"]
        
        # Find the paragraph containing this position
        char_count = 0
        target_paragraph = None
        
        for i, para in enumerate(doc_obj.paragraphs):
            if char_count <= position < char_count + len(para.text):
                relative_pos = position - char_count
                # Split text at issue position
                text = para.text
                para.text = ""
                
                # First part
                if relative_pos > 0:
                    para.add_run(text[:relative_pos])
                
                # Flagged part with red color and comment
                issue_text_end = min(relative_pos + len(issue["matched_text"]), len(text))
                flagged_run = para.add_run(text[relative_pos:issue_text_end])
                flagged_run.font.color.rgb = 0xFF0000  # Red color
                
                # Since python-docx doesn't directly support comments, we'll use a visual marker
                comment_text = f" [ISSUE: {issue['issue']}. {issue['suggestion']}]"
                comment_run = para.add_run(comment_text)
                comment_run.font.color.rgb = 0x0000FF  # Blue color
                comment_run.font.italic = True
                
                # Rest of paragraph
                if issue_text_end < len(text):
                    para.add_run(text[issue_text_end:])
                
                break
            
            char_count += len(para.text) + 1  # +1 for newline character
    
    return doc_obj

def get_binary_file_downloader_html(bin_file, file_label='File'):
    """Generate a download link for a binary file"""
    b64 = base64.b64encode(bin_file.getvalue()).decode()
    return f'<a href="data:application/octet-stream;base64,{b64}" download="{file_label}.docx">Download {file_label}</a>'

def process_uploaded_documents(uploaded_files):
    """Process all uploaded documents"""
    document_classifier = create_document_type_classifier()
    all_documents = []
    all_issues = []
    recognized_documents = []
    
    for uploaded_file in uploaded_files:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(uploaded_file.getbuffer())
            tmp_file_path = tmp_file.name
        
        try:
            # Read the document
            doc = Document(tmp_file_path)
            
            # Extract full text for classification
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Classify the document
            document_type = classify_document(full_text, document_classifier)
            if document_type:
                category = next((info['category'] for doc, info in document_classifier.items() 
                                if doc == document_type), None)
            else:
                document_type = "Unknown Document"
                category = "Unknown"
            
            # Detect red flags
            issues = detect_red_flags(full_text, document_type)
            
            # Add comments to document
            marked_doc = insert_comments(doc, issues)
            
            # Save the modified document for download
            output_file = BytesIO()
            marked_doc.save(output_file)
            output_file.seek(0)
            
            all_documents.append({
                "name": uploaded_file.name,
                "type": document_type,
                "category": category,
                "modified_file": output_file,
                "issues_count": len(issues)
            })
            
            all_issues.extend(issues)
            recognized_documents.append(document_type)
            
        except Exception as e:
            st.error(f"Error processing {uploaded_file.name}: {str(e)}")
        
        finally:
            # Clean up temp file
            if os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)
    
    return all_documents, all_issues, recognized_documents

def determine_process_type(recognized_documents):
    """Determine what legal process the user is attempting based on uploaded documents"""
    process_match = {}
    
    for category, required_docs in DOCUMENT_REQUIREMENTS.items():
        docs_found = [doc for doc in recognized_documents if doc in required_docs]
        match_percentage = len(docs_found) / len(required_docs) if required_docs else 0
        process_match[category] = {
            "match_percentage": match_percentage,
            "docs_found": docs_found,
            "required_docs": required_docs,
            "missing_docs": [doc for doc in required_docs if doc not in docs_found]
        }
    
    # Find the best matching process
    best_match = max(process_match.items(), key=lambda x: x[1]["match_percentage"])
    
    if best_match[1]["match_percentage"] > 0.3:  # At least 30% match to consider it valid
        return best_match[0], best_match[1]
    else:
        return "Unknown", {"match_percentage": 0, "docs_found": [], "required_docs": [], "missing_docs": []}

def find_related_resources(document_type):
    """Find official ADGM resources related to the document type"""
    related = []
    
    # Convert document_type to lowercase for case-insensitive matching
    doc_lower = document_type.lower()
    
    for resource in ADGM_RESOURCES:
        # Extract keywords from the document type
        keywords = re.findall(r'\b[A-Za-z]{3,}\b', document_type)
        keywords = [k.lower() for k in keywords if k.lower() not in ['and', 'the', 'for', 'of']]
        
        # Check if any keyword matches the resource category or document type
        category_match = any(keyword in resource['category'].lower() for keyword in keywords)
        type_match = any(keyword in resource['document_type'].lower() for keyword in keywords)
        
        if category_match or type_match:
            related.append(resource)
    
    return related

def main():
    # App header
    st.title("ADGM-Compliant Corporate Agent")
    st.markdown("""
    This intelligent agent reviews and validates documentation for business incorporation and compliance 
    within the Abu Dhabi Global Market (ADGM) jurisdiction.
    """)
    
    # Main navigation tabs
    tabs = st.tabs(["Document Review", "Official Resources", "About ADGM"])
    
    # Document Review Tab
    with tabs[0]:
        # File upload section
        st.header("Document Upload")
        uploaded_files = st.file_uploader("Upload ADGM-related documents (.docx format)", 
                                        type=["docx"], accept_multiple_files=True)
        
        if not uploaded_files:
            st.info("Please upload .docx files to begin the review process.")
            st.markdown("""
            ### Document Types Handled
            
            #### Company Formation Documents
            * Articles of Association (AoA)
            * Memorandum of Association (MoA/MoU)
            * Board Resolution Templates
            * Shareholder Resolution Templates
            * Incorporation Application Form
            * UBO Declaration Form
            * Register of Members and Directors
            * Change of Registered Address Notice
            
            #### Other Categories
            * Licensing Regulatory Filings
            * Employment HR Contracts
            * Commercial Agreements
            * Compliance Risk Policies
            """)
            return
        
        with st.spinner("Processing documents... This may take a moment."):
            all_documents, all_issues, recognized_documents = process_uploaded_documents(uploaded_files)
        
        # Process identification
        process_type, process_details = determine_process_type(recognized_documents)
        
        # Display process information
        st.header("Process Identification")
        if process_type != "Unknown":
            st.success(f"Identified Process: {process_type}")
            st.metric("Documents Found", f"{len(process_details['docs_found'])} of {len(process_details['required_docs'])}")
            
            # Display missing documents notification
            if process_details['missing_docs']:
                st.warning("Missing Documents:")
                for doc in process_details['missing_docs']:
                    st.markdown(f"* {doc}")
                
                st.markdown(f"""
                It appears that you're trying to complete a {process_type} process in ADGM.
                Based on our reference list, you have uploaded {len(process_details['docs_found'])} out of {len(process_details['required_docs'])} required documents.
                """)
        else:
            st.info("Unable to determine the specific legal process. Please ensure you've uploaded the relevant documents.")
        
        # Document review section
        st.header("Document Review Results")
        
        if not all_documents:
            st.error("No valid documents were processed.")
            return
        
        # Summary of findings
        total_issues = len(all_issues)
        high_severity_issues = len([issue for issue in all_issues if issue['severity'] == 'High'])
        medium_severity_issues = len([issue for issue in all_issues if issue['severity'] == 'Medium'])
        low_severity_issues = len([issue for issue in all_issues if issue['severity'] == 'Low'])
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Issues", total_issues)
        with col2:
            st.metric("High Severity", high_severity_issues, delta=None, delta_color="inverse")
        with col3:
            st.metric("Medium Severity", medium_severity_issues, delta=None, delta_color="inverse")
        with col4:
            st.metric("Low Severity", low_severity_issues, delta=None, delta_color="off")
        
        # Document-specific results
        st.subheader("Documents Analyzed")
        for idx, doc in enumerate(all_documents):
            with st.expander(f"{doc['name']} ({doc['type']}): {doc['issues_count']} issues"):
                st.markdown(f"**Document Type**: {doc['type']}")
                st.markdown(f"**Category**: {doc['category']}")
                
                # Generate download link for modified document
                st.markdown(get_binary_file_downloader_html(doc['modified_file'], 
                                                        f"Reviewed_{doc['name']}"), unsafe_allow_html=True)
                
                # Show issues for this document
                doc_issues = [issue for issue in all_issues if issue['document'] == doc['type']]
                if doc_issues:
                    st.markdown("### Issues Found:")
                    for issue_idx, issue in enumerate(doc_issues):
                        severity_color = {
                            "High": "🔴",
                            "Medium": "🟠",
                            "Low": "🟡"
                        }.get(issue['severity'], "⚪")
                        
                        st.markdown(f"{severity_color} **{issue['severity']}**: {issue['issue']}")
                        st.markdown(f"**Suggestion**: {issue['suggestion']}")
                        st.markdown(f"**Found in**: \"{issue['matched_text']}\"")
                        st.markdown("---")
                else:
                    st.success("No issues found in this document.")
                
                # Show related official resources
                related_resources = find_related_resources(doc['type'])
                if related_resources:
                    st.markdown("### Related Official ADGM Resources:")
                    for resource in related_resources:
                        st.markdown(f"- [{resource['document_type']}]({resource['link']}) - *{resource['category']}*")
        
        # Generate structured report
        if process_type != "Unknown":
            report = {
                "process": process_type,
                "documents_uploaded": len(process_details['docs_found']),
                "required_documents": len(process_details['required_docs']),
                "missing_documents": process_details['missing_docs'],
                "issues_found": [{
                    "document": issue["document"],
                    "section": issue["section"],
                    "issue": issue["issue"],
                    "severity": issue["severity"],
                    "suggestion": issue["suggestion"]
                } for issue in all_issues]
            }
            
            st.header("Structured Report")
            st.json(report)
            
            # Download report option
            report_json = json.dumps(report, indent=2)
            st.download_button(
                label="Download JSON Report",
                data=report_json,
                file_name=f"ADGM_Compliance_Report_{process_type.replace(' ', '_')}.json",
                mime="application/json"
            )
    
    # Official Resources Tab
    with tabs[1]:
        st.header("Official ADGM Document Resources")
        st.markdown("""
        These links provide direct access to official ADGM templates, forms, and guidance documents that will help ensure your documents 
        are compliant with the latest ADGM requirements.
        """)
        
        # Group resources by category
        categories = {}
        for resource in ADGM_RESOURCES:
            if resource['category'] not in categories:
                categories[resource['category']] = []
            categories[resource['category']].append(resource)
        
        # Display resources by category
        for category, resources in categories.items():
            with st.expander(f"{category} ({len(resources)} resources)"):
                for resource in resources:
                    st.markdown(f"- [{resource['document_type']}]({resource['link']})")
    
    # About ADGM Tab
    with tabs[2]:
        st.header("About ADGM")
        st.markdown("""
        ### Abu Dhabi Global Market (ADGM)
        
        ADGM is an international financial center and free zone established in Abu Dhabi, UAE. It offers its own legal framework 
        based on English common law, separate from the UAE federal laws, creating a business-friendly environment for international 
        companies.
        
        ### Key Features of ADGM
        
        - **Legal Framework**: ADGM has its own civil and commercial laws based on English common law
        - **Independent Courts**: ADGM Courts provide an independent judiciary
        - **Regulatory Framework**: Financial services are regulated by the Financial Services Regulatory Authority (FSRA)
        - **Company Types**: Various legal structures including LLCs, SPVs, branches, and foundations
        - **100% Foreign Ownership**: Allows complete foreign ownership of businesses
        
        ### Benefits of ADGM Compliance
        
        - Access to a stable financial center with strong regulatory oversight
        - Protection through a legal system based on English common law
        - Strategic location connecting East and West markets
        - Tax efficiency with 0% corporate and personal income tax
        - Access to a growing ecosystem of financial and professional services
        
        For more information, visit the [official ADGM website](https://www.adgm.com/).
        """)

if __name__ == "__main__":
    main()