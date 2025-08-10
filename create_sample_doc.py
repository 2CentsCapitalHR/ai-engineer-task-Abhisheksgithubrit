from docx import Document
from docx.shared import Pt

# Create a new Document
doc = Document()

# Add a title
title = doc.add_heading('ARTICLES OF ASSOCIATION', level=0)
title.alignment = 1  # Center alignment

# Add subtitle
subtitle = doc.add_heading('OF', level=1)
subtitle.alignment = 1
company_name = doc.add_heading('EXAMPLE TRADING LTD', level=1)
company_name.alignment = 1

doc.add_paragraph()  # Add some space

# Add sections with issues that should be detected
sections = [
    {
        'title': '1. INTERPRETATION',
        'content': 'In these Articles, unless the context otherwise requires:\n'
                 '"Act" means the UAE Federal Law No. 2 of 2015 concerning Commercial Companies;\n'
                 '"AED" means United Arab Emirates Dirhams, the official currency of the UAE;\n'
                 '"Board" means the board of directors of the Company;\n'
                 '"Company" means Example Trading Ltd;\n'
                 '"Director" means a director of the Company;\n'
                 '"Ordinary Resolution" means a resolution passed by a simple majority of votes cast;\n'
                 '"Special Resolution" means a resolution passed by not less than three-fourths of votes cast;\n'
                 '"UAE" means the United Arab Emirates.'
    },
    {
        'title': '2. REGISTERED OFFICE',
        'content': 'The registered office of the Company will be situated in Abu Dhabi Global Market, Abu Dhabi, UAE.'
    },
    {
        'title': '3. LIABILITY OF MEMBERS',
        'content': 'The liability of the members is limited to the amount, if any, unpaid on the shares held by them.'
    },
    {
        'title': '4. SHARE CAPITAL',
        'content': 'The share capital of the Company is AED 1,000,000 divided into 1,000,000 shares of AED 1 each.'
    },
    {
        'title': '5. DIRECTORS',
        'content': 'The Company shall have a minimum of three directors. At least one director must be a natural person.'
    },
    {
        'title': '6. PROCEEDINGS OF DIRECTORS',
        'content': '6.1 The Directors may meet together for the dispatch of business, adjourn and otherwise regulate their meetings as they think fit.\n'
                 '6.2 Questions arising at any meeting shall be decided by a majority of votes. In case of an equality of votes, the chairman shall have a second or casting vote.'
    },
    {
        'title': '7. GENERAL MEETINGS',
        'content': '7.1 The Company shall hold an annual general meeting each year.\n'
                 '7.2 A general meeting may be called by the Directors whenever they think fit.\n'
                 '7.3 General meetings shall be held at the registered office of the Company or at such other place as may be determined by the Directors.'
    },
    {
        'title': '8. NOTICES',
        'content': '8.1 A notice may be given by the Company to any member either personally or by sending it by post to him at his registered address.\n'
                 '8.2 Where a notice is sent by post, service of the notice shall be deemed to be effected by properly addressing, prepaying, and posting a letter containing the notice.'
    },
    {
        'title': '9. INDEMNITY',
        'content': 'Every Director shall be indemnified out of the assets of the Company against any liability incurred by him in defending any proceedings in relation to the Company.'
    },
    {
        'title': '10. DISPUTES',
        'content': 'Any disputes arising out of or in connection with these Articles shall be subject to the exclusive jurisdiction of the Dubai Courts.'
    }
]

# Add each section to the document
for section in sections:
    heading = doc.add_heading(section['title'], level=2)
    paragraph = doc.add_paragraph(section['content'])
    
    # Add some spacing after each section
    doc.add_paragraph()

# Save the document
doc.save('sample_articles_of_association.docx')

print("Sample Articles of Association document created successfully.")