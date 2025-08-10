from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add a title
title = doc.add_heading('UBO DECLARATION FORM', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('ADGM Registration Authority')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('COMPANY DETAILS')
p = doc.add_paragraph()
p.add_run('Company Name: ').bold = True
p.add_run('Example Trading Ltd')
p = doc.add_paragraph()
p.add_run('License Number: ').bold = True
p.add_run('ADGM-C0123456')
p = doc.add_paragraph()
p.add_run('Address: ').bold = True
p.add_run('Al Maryah Island, Abu Dhabi, UAE')

doc.add_heading('BENEFICIAL OWNER DETAILS', level=1)

# Create a table for UBO details
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
# Set header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Name'
header_cells[1].text = 'Nationality'
header_cells[2].text = 'Date of Birth'
header_cells[3].text = 'Ownership (%)'
header_cells[4].text = 'Type of Control'

# Add data rows with intentional issues
data = [
    ('John Smith', 'British', '15/04/1980', '15%', 'Direct Shareholder'),
    ('Sarah Johnson', 'American', '22/06/1975', '20%', 'Indirect Shareholder'),
    ('Ali Ahmed', 'Emirati', '10/11/1982', '10%', 'Board Control')
]

for name, nationality, dob, ownership, control in data:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = nationality
    row_cells[2].text = dob
    row_cells[3].text = ownership
    row_cells[4].text = control

# Add declaration section with an issue (incorrect threshold)
doc.add_heading('DECLARATION', level=1)
declaration = doc.add_paragraph()
declaration.add_run('I declare that:').bold = True
doc.add_paragraph('1. I am authorized to submit this declaration on behalf of the company.')
doc.add_paragraph('2. The information provided in this form is complete and accurate.')
doc.add_paragraph('3. The company has disclosed all beneficial owners who own or control 15% or more of the company shares or voting rights.')
doc.add_paragraph('4. The company will notify the registrar of any changes to the information provided within 15 days of such change.')

doc.add_paragraph()
signature_line = doc.add_paragraph('________________')
signature_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
name_line = doc.add_paragraph('Name and Position')
name_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
date_line = doc.add_paragraph('Date: ___/___/_____')
date_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Add legal reference with UAE federal reference instead of ADGM
doc.add_paragraph()
legal_ref = doc.add_paragraph('This declaration is made pursuant to the UAE Federal Ultimate Beneficial Owner Regulations.')
legal_ref.style = 'Caption'

# Save the document
doc.save('sample_ubo_declaration.docx')

print("Sample UBO Declaration document created successfully.")